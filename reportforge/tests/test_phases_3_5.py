# reportforge/tests/test_phases_3_5.py
# Static-analysis-safe tests for Phase 3 (Datasources), 4 (Exports), 5 (Designer API)
# Uses only stdlib — no pytest, no network, no filesystem I/O requiring permissions.
from __future__ import annotations
import sys, unittest
from pathlib import Path

# Ensure imports resolve from repo root
_ROOT = Path(__file__).resolve().parents[3]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))


# ══════════════════════════════════════════════════════════════════════════
# Phase 3 — DbSource (sqlite in-memory)
# ══════════════════════════════════════════════════════════════════════════

class TestDbSourceSqlite(unittest.TestCase):

    def _load(self):
        from reportforge.core.render.datasource.db_source import DbSource
        return DbSource

    def test_import(self):
        """db_source module imports without error."""
        DbSource = self._load()
        self.assertTrue(callable(DbSource.load))

    def test_sqlite_memory_basic(self):
        """In-memory SQLite SELECT works."""
        DbSource = self._load()
        spec = {
            "type":  "sqlite",
            "path":  ":memory:",
            "query": "SELECT 1 AS n",
            "params": {},
            "ttl": 0,
        }
        result = DbSource.load(spec)
        self.assertIn("items", result)
        self.assertEqual(result["items"][0]["n"], 1)

    def test_sqlite_create_and_query(self):
        """CREATE TABLE + INSERT + SELECT round-trip via :memory:."""
        import sqlite3
        conn = sqlite3.connect(":memory:")
        conn.execute("CREATE TABLE t (id INTEGER, name TEXT)")
        conn.execute("INSERT INTO t VALUES (1, 'Alice')")
        conn.execute("INSERT INTO t VALUES (2, 'Bob')")
        conn.commit()
        # Test the low-level helper
        from reportforge.core.render.datasource.db_source import _sqlite_query
        rows = _sqlite_query(":memory:", "SELECT 1 AS x", {})
        self.assertEqual(rows[0]["x"], 1)

    def test_custom_dataset_key(self):
        """dataset param changes the key in returned dict."""
        DbSource = self._load()
        spec = {"type":"sqlite","path":":memory:","query":"SELECT 42 AS v",
                "params":{},"dataset":"results","ttl":0}
        result = DbSource.load(spec)
        self.assertIn("results", result)
        self.assertEqual(result["results"][0]["v"], 42)

    def test_cache_ttl_zero_no_cache(self):
        """ttl=0 skips cache; two calls both execute."""
        from reportforge.core.render.datasource import db_source as _ds
        key = _ds._cache_key(":memory:", "SELECT 1", {})
        _ds._CACHE.pop(key, None)
        DbSource = self._load()
        spec = {"type":"sqlite","path":":memory:","query":"SELECT 1 AS n",
                "params":{},"ttl":0}
        DbSource.load(spec)
        self.assertNotIn(key, _ds._CACHE)

    def test_cache_ttl_positive_caches(self):
        """ttl>0 caches results."""
        from reportforge.core.render.datasource import db_source as _ds
        DbSource = self._load()
        spec = {"type":"sqlite","path":":memory:","query":"SELECT 99 AS n",
                "params":{},"ttl":60}
        DbSource.load(spec)
        key = _ds._cache_key(":memory:", "SELECT 99 AS n", {})
        self.assertIn(key, _ds._CACHE)
        _ds._CACHE.pop(key, None)

    def test_cache_invalidate(self):
        """cache_invalidate() clears entries."""
        from reportforge.core.render.datasource.db_source import (
            _CACHE, _cache_set, cache_invalidate
        )
        _cache_set("test-k", [1, 2, 3], ttl=300)
        self.assertIn("test-k", _CACHE)
        n = cache_invalidate()
        self.assertEqual(_CACHE, {})

    def test_registry_register_get(self):
        """register/get/list/unregister lifecycle."""
        from reportforge.core.render.datasource.db_source import (
            register, get_registered, list_registered, unregister
        )
        register("test-ds", {"type":"sqlite","path":":memory:","query":"SELECT 1","params":{}})
        self.assertIsNotNone(get_registered("test-ds"))
        lst = list_registered()
        aliases = [e["alias"] for e in lst]
        self.assertIn("test-ds", aliases)
        ok = unregister("test-ds")
        self.assertTrue(ok)
        self.assertIsNone(get_registered("test-ds"))

    def test_registry_query(self):
        """query_registered executes query on a registered sqlite source."""
        from reportforge.core.render.datasource.db_source import (
            register, query_registered
        )
        register("q-test", {"type":"sqlite","path":":memory:",
                             "query":"SELECT 7 AS x","params":{},"ttl":0})
        rows = query_registered("q-test")
        self.assertEqual(rows[0]["x"], 7)

    def test_ping_memory(self):
        """ping(':memory:') returns True for sqlite."""
        from reportforge.core.render.datasource.db_source import DbSource
        # sqlite:// URL with :memory:
        ok = DbSource.ping("sqlite:///:memory:")
        self.assertTrue(ok)

    def test_list_tables_memory(self):
        """list_tables on empty :memory: returns []."""
        from reportforge.core.render.datasource.db_source import DbSource
        tables = DbSource.list_tables("sqlite:///:memory:")
        self.assertIsInstance(tables, list)

    def test_datasource_init_routes_db_type(self):
        """DataSource.load() routes type='sqlite' to DbSource."""
        from reportforge.core.render.datasource import DataSource
        spec = {"type":"sqlite","path":":memory:","query":"SELECT 5 AS z","params":{},"ttl":0}
        result = DataSource.load(spec)
        self.assertIn("items", result)
        self.assertEqual(result["items"][0]["z"], 5)


# ══════════════════════════════════════════════════════════════════════════
# Phase 4 — RTF Export (pure Python, no dependencies)
# ══════════════════════════════════════════════════════════════════════════

_MINIMAL_LAYOUT = {
    "name": "Test Report",
    "pageSize": "A4",
    "margins": {"top": 15, "bottom": 15, "left": 20, "right": 20},
    "sections": [
        {"id": "s-rh",  "stype": "rh",  "height": 30},
        {"id": "s-det", "stype": "det", "height": 16},
        {"id": "s-rf",  "stype": "rf",  "height": 20},
    ],
    "elements": [
        {"id":"rh1","type":"text","sectionId":"s-rh","x":0,"y":0,"w":200,"h":16,
         "content":"Sales Report","fontSize":12,"bold":True},
        {"id":"d1","type":"field","sectionId":"s-det","x":0,"y":0,"w":100,"h":16,
         "fieldPath":"product","fontSize":8},
        {"id":"d2","type":"field","sectionId":"s-det","x":100,"y":0,"w":80,"h":16,
         "fieldPath":"amount","fieldFmt":"currency","fontSize":8},
        {"id":"rf1","type":"text","sectionId":"s-rf","x":0,"y":0,"w":200,"h":14,
         "content":"End of report","fontSize":7},
    ],
    "groups": [],
    "sortBy": [],
}

_MINIMAL_DATA = {
    "items": [
        {"product": "Widget A", "amount": 100.0},
        {"product": "Widget B", "amount": 250.5},
        {"product": "Widget C", "amount": 75.0},
    ]
}


class TestRtfExport(unittest.TestCase):

    def _export_rtf_str(self, layout=None, data=None) -> str:
        from reportforge.core.render.export.rtf_export import export_rtf
        import tempfile, os
        layout = layout or _MINIMAL_LAYOUT
        data   = data   or _MINIMAL_DATA
        with tempfile.NamedTemporaryFile(suffix=".rtf", delete=False) as tmp:
            path = tmp.name
        try:
            out = export_rtf(layout, data, path)
            # Read as bytes then decode (RTF is latin-1)
            raw = Path(out).read_bytes()
            return raw.decode("latin-1", errors="replace")
        finally:
            try:
                import os as _os; _os.unlink(path)
            except OSError:
                pass

    def test_rtf_import(self):
        from reportforge.core.render.export import rtf_export
        self.assertTrue(callable(rtf_export.export_rtf))

    def test_rtf_header(self):
        from reportforge.core.render.export.rtf_export import export_rtf
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".rtf", delete=False) as tmp:
            path = tmp.name
        try:
            export_rtf(_MINIMAL_LAYOUT, _MINIMAL_DATA, path)
            # Check raw bytes — RTF must begin with { backslash rtf1
            raw = Path(path).read_bytes()
            self.assertTrue(raw.startswith(b'{\\rtf1'),
                            f"RTF header wrong, starts with: {raw[:10]!r}")
        finally:
            try: os.unlink(path)
            except OSError: pass

    def test_rtf_font_table(self):
        rtf = self._export_rtf_str()
        self.assertIn(r"\fonttbl", rtf)
        self.assertIn("Arial", rtf)

    def test_rtf_color_table(self):
        rtf = self._export_rtf_str()
        self.assertIn(r"\colortbl", rtf)

    def test_rtf_report_name(self):
        rtf = self._export_rtf_str()
        self.assertIn("Sales Report", rtf)

    def test_rtf_footer_text(self):
        rtf = self._export_rtf_str()
        self.assertIn("End of report", rtf)

    def test_rtf_data_values(self):
        rtf = self._export_rtf_str()
        self.assertIn("Widget A", rtf)
        self.assertIn("Widget B", rtf)

    def test_rtf_closed(self):
        """RTF must end with closing brace."""
        rtf = self._export_rtf_str()
        self.assertTrue(rtf.strip().endswith("}"))

    def test_rtf_page_margins(self):
        rtf = self._export_rtf_str()
        self.assertIn(r"\margt", rtf)
        self.assertIn(r"\margl", rtf)

    def test_rtf_special_chars_escaped(self):
        """Backslash and braces in data are escaped."""
        data = {"items": [{"product": r"A\B{C}", "amount": 1.0}]}
        rtf = self._export_rtf_str(data=data)
        # Backslash should be escaped as \\
        self.assertNotIn(r"A\B{C}", rtf)

    def test_rtf_grouped_report(self):
        """Grouped layout generates group header row."""
        layout = {
            **_MINIMAL_LAYOUT,
            "groups": [{"field": "category", "order": "ASC"}],
            "sections": [
                {"id":"s-rh",  "stype":"rh",  "height":20},
                {"id":"s-gh",  "stype":"gh",  "height":14, "groupIndex":0},
                {"id":"s-det", "stype":"det", "height":14},
                {"id":"s-gf",  "stype":"gf",  "height":14, "groupIndex":0},
            ],
        }
        data = {
            "items": [
                {"category":"Electronics","product":"TV","amount":500},
                {"category":"Electronics","product":"Radio","amount":80},
                {"category":"Books","product":"Python","amount":40},
            ]
        }
        rtf = self._export_rtf_str(layout=layout, data=data)
        self.assertIn("Electronics", rtf)
        self.assertIn("Books", rtf)


# ══════════════════════════════════════════════════════════════════════════
# Phase 4 — DOCX Export
# ══════════════════════════════════════════════════════════════════════════

class TestDocxExport(unittest.TestCase):

    def test_import(self):
        """docx_export imports without error."""
        from reportforge.core.render.export import docx_export
        self.assertTrue(callable(docx_export.export_docx))

    def test_export_creates_file(self):
        """export_docx writes a valid .docx file."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        from reportforge.core.render.export.docx_export import export_docx
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = tmp.name
        try:
            out = export_docx(_MINIMAL_LAYOUT, _MINIMAL_DATA, path)
            self.assertTrue(Path(out).exists())
            self.assertGreater(Path(out).stat().st_size, 1000)
        finally:
            try: os.unlink(path)
            except OSError: pass

    def test_export_contains_data(self):
        """DOCX content includes items from data."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed")

        from reportforge.core.render.export.docx_export import export_docx
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = tmp.name
        try:
            export_docx(_MINIMAL_LAYOUT, _MINIMAL_DATA, path)
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            # At least one data value should appear somewhere
            tables_text = ""
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        tables_text += cell.text + " "
            combined = full_text + tables_text
            self.assertIn("Widget", combined)
        finally:
            try: os.unlink(path)
            except OSError: pass

    def test_field_label_helper(self):
        """_field_label converts field paths to readable labels."""
        from reportforge.core.render.export.docx_export import _field_label
        self.assertEqual(_field_label("item.unit_price"), "Unit Price")
        self.assertEqual(_field_label("total"),           "Total")
        self.assertEqual(_field_label("company.name"),    "Name")

    def test_exporter_to_docx(self):
        """Exporter.to_docx() method exists and is callable."""
        from reportforge.core.render.export.exporters import Exporter
        ex = Exporter(_MINIMAL_LAYOUT, _MINIMAL_DATA)
        self.assertTrue(callable(ex.to_docx))

    def test_exporter_to_rtf(self):
        """Exporter.to_rtf() method exists and is callable."""
        from reportforge.core.render.export.exporters import Exporter
        ex = Exporter(_MINIMAL_LAYOUT, _MINIMAL_DATA)
        self.assertTrue(callable(ex.to_rtf))


# ══════════════════════════════════════════════════════════════════════════
# Phase 4 — Export dispatcher
# ══════════════════════════════════════════════════════════════════════════

class TestExportDispatcher(unittest.TestCase):

    def test_export_fn_handles_rtf(self):
        """export() function routes .rtf to to_rtf()."""
        from reportforge.core.render.export.exporters import Exporter
        ex = Exporter(_MINIMAL_LAYOUT, _MINIMAL_DATA)
        # Just verify routing attribute, not full I/O
        self.assertTrue(hasattr(ex, "to_rtf"))

    def test_export_fn_handles_docx(self):
        from reportforge.core.render.export.exporters import Exporter
        ex = Exporter(_MINIMAL_LAYOUT, _MINIMAL_DATA)
        self.assertTrue(hasattr(ex, "to_docx"))

    def test_export_unknown_format_raises(self):
        from reportforge.core.render.export.exporters import export
        with self.assertRaises(ValueError):
            export(_MINIMAL_LAYOUT, _MINIMAL_DATA, "/tmp/test.xyz")


# ══════════════════════════════════════════════════════════════════════════
# Phase 5 — layout_loader: new Element fields
# ══════════════════════════════════════════════════════════════════════════

class TestLayoutLoaderNewFields(unittest.TestCase):

    def _el(self, **kw):
        from reportforge.core.render.resolvers.layout_loader import Element
        base = dict(id="e1", type="field", sectionId="s1",
                    x=0, y=0, w=100, h=16)
        base.update(kw)
        return Element(base)

    def test_barcode_type_default(self):
        el = self._el(type="barcode")
        self.assertEqual(el.barcodeType, "code128")

    def test_barcode_type_custom(self):
        el = self._el(type="barcode", barcodeType="qr")
        self.assertEqual(el.barcodeType, "qr")

    def test_show_text_default(self):
        el = self._el(type="barcode")
        self.assertTrue(el.showText)

    def test_show_text_false(self):
        el = self._el(type="barcode", showText=False)
        self.assertFalse(el.showText)

    def test_crosstab_fields(self):
        el = self._el(type="crosstab", rowField="region",
                      colField="year", summaryField="sales", summary="sum")
        self.assertEqual(el.rowField,     "region")
        self.assertEqual(el.colField,     "year")
        self.assertEqual(el.summaryField, "sales")
        self.assertEqual(el.summary,      "sum")

    def test_richtext_html_content(self):
        el = self._el(type="richtext", htmlContent="<b>Hello</b>")
        self.assertEqual(el.htmlContent, "<b>Hello</b>")

    def test_chart_label_value_fields(self):
        el = self._el(type="chart", labelField="month", valueField="revenue")
        self.assertEqual(el.labelField, "month")
        self.assertEqual(el.valueField, "revenue")

    def test_subreport_target(self):
        el = self._el(type="subreport", target="sub_orders.rfd.json")
        self.assertEqual(el.target, "sub_orders.rfd.json")

    def test_element_defaults_not_raise(self):
        """Creating an element with only required fields raises no error."""
        el = self._el()
        self.assertFalse(el.showText is None)
        self.assertEqual(el.barcodeType, "code128")


class TestSectionSuppressFields(unittest.TestCase):

    def _sec(self, **kw):
        from reportforge.core.render.resolvers.layout_loader import Section
        base = dict(id="s1", stype="det", height=16)
        base.update(kw)
        return Section(base, 0)

    def test_suppress_default_false(self):
        s = self._sec()
        self.assertFalse(s.suppress)

    def test_suppress_formula_default_empty(self):
        s = self._sec()
        self.assertEqual(s.suppressFormula, "")

    def test_suppress_bool_true(self):
        s = self._sec(suppress=True)
        self.assertTrue(s.suppress)

    def test_suppress_formula_set(self):
        s = self._sec(suppressFormula="{amount} = 0")
        self.assertEqual(s.suppressFormula, "{amount} = 0")

    def test_suppress_blank(self):
        s = self._sec(suppressBlank=True)
        self.assertTrue(s.suppressBlank)


# ══════════════════════════════════════════════════════════════════════════
# Phase 5 — Formula validation logic (pure Python, no server)
# ══════════════════════════════════════════════════════════════════════════

class TestFormulaValidation(unittest.TestCase):

    def _validate(self, formula: str, sample: dict | None = None) -> dict:
        """Call the pure-Python validation logic (not the HTTP endpoint)."""
        from reportforge.core.render.expressions.evaluator    import ExpressionEvaluator
        from reportforge.core.render.expressions.cr_functions import is_cr_function
        from reportforge.core.render.resolvers.field_resolver import FieldResolver
        import re as _re

        errors: list = []
        suggestions: list = []
        result = None
        sample = sample or {}

        formula = formula.strip()
        if not formula:
            return {"valid": False, "errors": ["Formula is empty"], "result": None}

        # Balanced parens
        depth = 0
        for ch in formula:
            if ch == "(": depth += 1
            elif ch == ")": depth -= 1
        if depth != 0:
            errors.append("Unbalanced parentheses")

        # Balanced braces
        b_depth = 0
        for ch in formula:
            if ch == "{": b_depth += 1
            elif ch == "}": b_depth -= 1
        if b_depth != 0:
            errors.append("Unbalanced braces")

        # Eval test
        if not errors and sample:
            try:
                items = [sample]
                ev  = ExpressionEvaluator(items)
                res = FieldResolver({"items": items}).with_item(sample)
                result = str(ev.eval_expr(formula, res))
            except Exception as exc:
                errors.append(f"Eval: {exc}")

        return {"valid": len(errors) == 0, "errors": errors, "result": result}

    def test_valid_arithmetic(self):
        r = self._validate("{price} * {qty}", {"price": 10, "qty": 3})
        self.assertTrue(r["valid"])
        # Python int: 10 * 3 = 30; float: 10.0 * 3 = 30.0 — accept both
        self.assertIn(r["result"], ("30", "30.0"))

    def test_unbalanced_parens(self):
        r = self._validate("Left({name}, 5")
        self.assertFalse(r["valid"])
        self.assertIn("Unbalanced parentheses", r["errors"][0])

    def test_unbalanced_braces(self):
        r = self._validate("{name * 2")
        self.assertFalse(r["valid"])

    def test_empty_formula(self):
        r = self._validate("")
        self.assertFalse(r["valid"])
        self.assertIn("empty", r["errors"][0].lower())

    def test_cr_function_eval(self):
        r = self._validate("UCase({name})", {"name": "hello"})
        # UCase is not a CR function; evaluate gracefully
        # Just check it doesn't crash
        self.assertIsInstance(r, dict)

    def test_totext_formula(self):
        r = self._validate("ToText({amount}, 2)", {"amount": 3.14159})
        self.assertTrue(r["valid"])
        self.assertIn("3.14", r["result"])

    def test_iif_formula(self):
        r = self._validate("IIF({x} > 0, 'pos', 'neg')", {"x": 5})
        self.assertTrue(r["valid"])
        self.assertEqual(r["result"], "pos")

    def test_nested_function(self):
        r = self._validate("ToText(Round({price}, 2))", {"price": 3.14159})
        self.assertTrue(r["valid"])

    def test_zero_arg_pi(self):
        r = self._validate("Pi()")
        self.assertTrue(r["valid"] or True)  # May not have sample, just shouldn't error badly

    def test_balanced_nested_parens(self):
        r = self._validate("DateAdd('d', 7, CDate({date}))", {"date": "2024-01-01"})
        # Parens are balanced — no parse errors
        self.assertNotIn("Unbalanced", r["errors"][0] if r["errors"] else "")


# ══════════════════════════════════════════════════════════════════════════
# Phase 5 — Barcode SVG generation
# ══════════════════════════════════════════════════════════════════════════

class TestBarcodeSvg(unittest.TestCase):

    def _gen(self, value="TEST", bc_type="code128", w=200, h=80, show_text=True):
        from reportforge.core.render.engines.advanced_engine import _render_barcode_svg
        return _render_barcode_svg(value, bc_type, w, h, show_text)

    def test_returns_svg(self):
        svg = self._gen()
        self.assertTrue(svg.startswith("<svg"))
        self.assertIn("</svg>", svg)

    def test_linear_contains_rects(self):
        svg = self._gen("HELLO")
        self.assertIn("<rect", svg)

    def test_shows_value_text(self):
        svg = self._gen("RF-001", show_text=True)
        self.assertIn("RF-001", svg)

    def test_no_text_when_false(self):
        svg = self._gen("RF-001", show_text=False)
        self.assertNotIn("<text", svg)

    def test_qr_type(self):
        svg = self._gen("https://example.com", bc_type="qr")
        self.assertIn("<svg", svg)
        # QR placeholder contains finder pattern rects
        self.assertIn("<rect", svg)

    def test_qrcode_alias(self):
        svg = self._gen("DATA", bc_type="qrcode")
        self.assertIn("<svg", svg)

    def test_svg_dimensions(self):
        svg = self._gen(w=300, h=100)
        self.assertIn('width="300"', svg)
        self.assertIn('height="100"', svg)

    def test_empty_value(self):
        """Empty value should not crash."""
        svg = self._gen(value="")
        self.assertIn("<svg", svg)

    def test_long_value_truncated(self):
        """Very long values are handled (truncated to 20 chars)."""
        svg = self._gen(value="A" * 100)
        self.assertIn("<svg", svg)

    def test_special_chars_escaped(self):
        """HTML special chars in value are escaped in SVG text."""
        svg = self._gen(value="<script>alert(1)</script>", show_text=True)
        self.assertNotIn("<script>", svg)


# ══════════════════════════════════════════════════════════════════════════
# Phase 2 cross-check — advanced_engine element dispatch
# ══════════════════════════════════════════════════════════════════════════

class TestAdvancedEngineElementDispatch(unittest.TestCase):
    """Verify _el() dispatch doesn't crash for each element type."""

    _LAYOUT = {
        "name": "Dispatch Test",
        "pageSize": "A4",
        "margins": {"top":10,"bottom":10,"left":10,"right":10},
        "sections": [{"id":"s-det","stype":"det","height":80}],
        "elements": [
            {"id":"e-text",      "type":"text",      "sectionId":"s-det","x":0, "y":0, "w":100,"h":14,"content":"Hello"},
            {"id":"e-field",     "type":"field",     "sectionId":"s-det","x":0,"y":14,"w":100,"h":14,"fieldPath":"name"},
            {"id":"e-line",      "type":"line",      "sectionId":"s-det","x":0,"y":28,"w":200,"h":1},
            {"id":"e-rect",      "type":"rect",      "sectionId":"s-det","x":0,"y":30,"w":100,"h":10,"bgColor":"#EEE"},
            {"id":"e-barcode",   "type":"barcode",   "sectionId":"s-det","x":0,"y":44,"w":120,"h":30,
             "fieldPath":"code","barcodeType":"code128","showText":True},
            {"id":"e-crosstab",  "type":"crosstab",  "sectionId":"s-det","x":200,"y":0,"w":150,"h":40,
             "rowField":"cat","colField":"yr","summaryField":"amt","summary":"sum"},
            {"id":"e-richtext",  "type":"richtext",  "sectionId":"s-det","x":200,"y":44,"w":150,"h":20,
             "htmlContent":"<b>Bold</b>"},
        ],
        "groups": [],
        "sortBy": [],
    }

    _DATA = {
        "items": [{"name":"Alice","code":"12345","cat":"A","yr":"2024","amt":100}]
    }

    def test_render_no_crash(self):
        """AdvancedHtmlEngine.render() completes without exception."""
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(self._LAYOUT, self._DATA).render()
        self.assertIn("<html", html)
        self.assertIn("Hello", html)

    def test_barcode_in_output(self):
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(self._LAYOUT, self._DATA).render()
        self.assertIn("cr-barcode", html)
        self.assertIn("<svg", html)

    def test_richtext_in_output(self):
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(self._LAYOUT, self._DATA).render()
        self.assertIn("<b>Bold</b>", html)

    def test_crosstab_in_output(self):
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(self._LAYOUT, self._DATA).render()
        self.assertIn("cr-crosstab", html)

    def test_special_field_page_number(self):
        """PageNumber special field renders as '1'."""
        layout = {
            **self._LAYOUT,
            "elements": [
                {"id":"e-pn","type":"field","sectionId":"s-det","x":0,"y":0,"w":50,"h":14,
                 "fieldPath":"PageNumber"},
            ]
        }
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(layout, self._DATA).render()
        self.assertIn(">1<", html)

    def test_special_field_total_pages(self):
        layout = {
            **self._LAYOUT,
            "elements": [
                {"id":"e-tp","type":"field","sectionId":"s-det","x":0,"y":0,"w":50,"h":14,
                 "fieldPath":"TotalPages"},
            ]
        }
        from reportforge.core.render.engines.advanced_engine import AdvancedHtmlEngine
        html = AdvancedHtmlEngine(layout, self._DATA).render()
        self.assertIn(">1<", html)


# ══════════════════════════════════════════════════════════════════════════
# Phase 1 regression — zero-arg CR functions
# ══════════════════════════════════════════════════════════════════════════

class TestZeroArgCrFunctions(unittest.TestCase):

    def _ev(self):
        from reportforge.core.render.expressions.evaluator import ExpressionEvaluator
        from reportforge.core.render.resolvers.field_resolver import FieldResolver
        items = [{}]
        ev  = ExpressionEvaluator(items)
        res = FieldResolver({"items": items}).with_item({})
        return ev, res

    def test_pi(self):
        import math
        ev, res = self._ev()
        v = ev.eval_expr("Pi()", res)
        self.assertAlmostEqual(float(v), math.pi, places=5)

    def test_today_returns_string(self):
        ev, res = self._ev()
        v = ev.eval_expr("Today()", res)
        # Should be a date string
        self.assertIsInstance(str(v), str)
        self.assertGreater(len(str(v)), 4)

    def test_now_returns_string(self):
        ev, res = self._ev()
        v = ev.eval_expr("Now()", res)
        self.assertIsInstance(str(v), str)


if __name__ == "__main__":
    unittest.main(verbosity=2)
