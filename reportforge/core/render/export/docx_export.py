# core/render/export/docx_export.py
# DOCX export — Phase 4
# Maps ReportForge layout sections/elements to python-docx paragraphs,
# tables, images and headers/footers.
from __future__ import annotations
from pathlib import Path
from typing import Any


def export_docx(layout_raw: dict, data: dict,
                output_path: str | Path) -> Path:
    """
    Render a ReportForge layout to a .docx file.

    Mapping:
      rh/rf  → body paragraphs (report header / footer)
      ph/pf  → Word Header / Footer
      det    → table rows (one row per item)
      gh/gf  → bold group header / footer rows
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("DOCX export requires python-docx: pip install python-docx")

    from ..pipeline.normalizer import normalize_layout
    from ..resolvers.field_resolver import FieldResolver, format_value
    from ..expressions.evaluator import ExpressionEvaluator
    from ..expressions.aggregator import Aggregator

    norm     = normalize_layout(layout_raw)
    resolver = FieldResolver(data)
    items    = list(data.get("items", []))
    ev       = ExpressionEvaluator(items)
    agg      = Aggregator(items)

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────
    margins = norm.get("margins", {"top": 15, "bottom": 15, "left": 20, "right": 20})
    section = doc.sections[0]
    section.top_margin    = Cm(margins.get("top",    15) / 10)
    section.bottom_margin = Cm(margins.get("bottom", 15) / 10)
    section.left_margin   = Cm(margins.get("left",   20) / 10)
    section.right_margin  = Cm(margins.get("right",  20) / 10)

    # ── Helper: parse color ───────────────────────────────────────
    def _rgb(hex_color: str) -> RGBColor | None:
        c = (hex_color or "").lstrip("#")
        if len(c) == 6:
            try:
                return RGBColor(int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))
            except ValueError:
                pass
        return None

    _ALIGN = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def _apply_run(run, el: dict) -> None:
        run.bold      = el.get("bold", False)
        run.italic    = el.get("italic", False)
        run.underline = el.get("underline", False)
        run.font.name = el.get("fontFamily", "Arial")
        run.font.size = Pt(el.get("fontSize", 8))
        c = _rgb(el.get("color", "#000000"))
        if c:
            run.font.color.rgb = c

    def _resolve_el(el: dict, item: dict | None = None, gitems: list | None = None) -> str:
        """Resolve an element's display value."""
        tp = el.get("type", "text")
        if tp == "text":
            content = el.get("content", "")
            if ev.contains_expr(content):
                res = resolver.with_item(item) if item else resolver
                content = ev.eval_text(content, res)
            return content
        if tp == "field":
            fp  = el.get("fieldPath", "")
            fmt = el.get("fieldFmt")
            if not fp:
                return ""
            res = resolver.with_item(item) if item else resolver
            # Check special fields
            _SF = {"PageNumber": "1", "TotalPages": "?",
                   "PrintDate": __import__("datetime").date.today().strftime("%d/%m/%Y")}
            if fp in _SF:
                return _SF[fp]
            return res.get_formatted(fp, fmt)
        return ""

    def _add_section_para(sec_elements: list, item: dict | None = None,
                          gitems: list | None = None, bold_override: bool = False) -> None:
        """Add a paragraph for each text/field element in a section."""
        texts = []
        for el in sec_elements:
            if el.get("type") in ("text", "field"):
                texts.append((el, _resolve_el(el, item, gitems)))
        if not texts:
            return
        # Group elements into one paragraph (left-to-right by x position)
        texts.sort(key=lambda t: t[0].get("x", 0))
        combined = "  ".join(v for _, v in texts if v)
        if not combined.strip():
            return
        # Pick style from first text element
        lead_el = texts[0][0]
        p = doc.add_paragraph()
        p.alignment = _ALIGN.get(lead_el.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(combined)
        _apply_run(run, lead_el)
        if bold_override:
            run.bold = True

    # ── Build sections index ──────────────────────────────────────
    sec_map: dict[str, list] = {}  # sectionId → elements
    for el in norm.get("elements", []):
        sec_map.setdefault(el.get("sectionId", ""), []).append(el)

    secs_by_type: dict[str, list] = {}
    for s in norm.get("sections", []):
        secs_by_type.setdefault(s["stype"], []).append(s)

    def _els(stype: str) -> list[dict]:
        result = []
        for s in secs_by_type.get(stype, []):
            result.extend(sec_map.get(s["id"], []))
        return result

    # ── Report Header ─────────────────────────────────────────────
    rh_els = _els("rh")
    if rh_els:
        _add_section_para(rh_els, bold_override=False)

    # ── Page Header as Word header ────────────────────────────────
    ph_els = _els("ph")
    if ph_els:
        hdr = doc.sections[0].header
        hdr.is_linked_to_previous = False
        ph_texts = [(el, _resolve_el(el)) for el in ph_els if el.get("type") in ("text","field")]
        ph_texts.sort(key=lambda t: t[0].get("x", 0))
        combined = "  ".join(v for _, v in ph_texts if v)
        if combined.strip():
            hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            run = hp.add_run(combined)
            run.bold = True
            run.font.size = Pt(9)

    # ── Group detection ───────────────────────────────────────────
    groups = norm.get("groups", [])
    det_secs = secs_by_type.get("det", [])
    gh_secs  = secs_by_type.get("gh", [])
    gf_secs  = secs_by_type.get("gf", [])

    det_els_flat = [el for s in det_secs for el in sec_map.get(s["id"], [])
                    if el.get("type") in ("text","field") and el.get("fieldPath")]

    # ── Detail table ──────────────────────────────────────────────
    if det_els_flat:
        # Sort columns by x position
        cols = sorted(det_els_flat, key=lambda e: e.get("x", 0))
        n_cols = len(cols)

        tbl = doc.add_table(rows=1, cols=n_cols)
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        from docx.oxml.ns import qn as _qn
        for ci, col_el in enumerate(cols):
            label = col_el.get("content") or _field_label(col_el.get("fieldPath",""))
            cell = hdr_cells[ci]
            cell.text = label
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(label)
            run.bold = True
            run.font.size = Pt(8)
            # Header cell shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(_qn("w:val"), "clear")
            shd.set(_qn("w:color"), "auto")
            shd.set(_qn("w:fill"), "1A3A6B")
            tcPr.append(shd)
            run2 = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run("")
            run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Group iteration or flat items
        def _write_items(item_list: list, gitems: list | None = None) -> None:
            item_agg = Aggregator(gitems or item_list)
            for item in item_list:
                row_cells = tbl.add_row().cells
                item_res  = resolver.with_item(item)
                for ci, col_el in enumerate(cols):
                    val = _resolve_el(col_el, item)
                    row_cells[ci].text = str(val)
                    p = row_cells[ci].paragraphs[0]
                    p.alignment = _ALIGN.get(col_el.get("align","left"), WD_ALIGN_PARAGRAPH.LEFT)
                    if p.runs:
                        p.runs[0].font.size = Pt(8)

        if not groups:
            _write_items(items)
        else:
            from itertools import groupby as _groupby
            gf_field = groups[0]["field"]
            sorted_items = sorted(items, key=lambda it: str(it.get(gf_field,"")))
            for gval, git in _groupby(sorted_items, key=lambda it: it.get(gf_field,"")):
                gitems = list(git)
                # Group header row
                gh_els = [el for s in gh_secs for el in sec_map.get(s["id"],[])
                          if el.get("type") in ("text","field")]
                if gh_els:
                    gr = tbl.add_row()
                    merged = gr.cells[0].merge(gr.cells[-1])
                    merged.text = str(gval)
                    if merged.paragraphs[0].runs:
                        merged.paragraphs[0].runs[0].bold = True
                        merged.paragraphs[0].runs[0].font.size = Pt(8)
                _write_items(gitems, gitems)
                # Group footer row
                gf_els = [el for s in gf_secs for el in sec_map.get(s["id"],[])
                          if el.get("type") in ("text","field")]
                if gf_els:
                    gfr = tbl.add_row()
                    merged = gfr.cells[0].merge(gfr.cells[-1])
                    vals = [_resolve_el(e, gitems[-1], gitems) for e in gf_els if e.get("type") in ("text","field")]
                    merged.text = "  ".join(v for v in vals if v)
                    if merged.paragraphs[0].runs:
                        merged.paragraphs[0].runs[0].bold = True
                        merged.paragraphs[0].runs[0].font.size = Pt(8)

    # ── Report Footer ─────────────────────────────────────────────
    rf_els = _els("rf")
    if rf_els:
        doc.add_paragraph()
        _add_section_para(rf_els)

    # ── Page Footer as Word footer ────────────────────────────────
    pf_els = _els("pf")
    if pf_els:
        ftr = doc.sections[0].footer
        ftr.is_linked_to_previous = False
        pf_texts = [(el, _resolve_el(el)) for el in pf_els if el.get("type") in ("text","field")]
        pf_texts.sort(key=lambda t: t[0].get("x", 0))
        combined = "  ".join(v for _, v in pf_texts if v)
        if combined.strip():
            fp2 = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp2.add_run(combined).font.size = Pt(7)

    # ── Save ──────────────────────────────────────────────────────
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _field_label(fp: str) -> str:
    part = fp.split(".")[-1] if "." in fp else fp
    return part.replace("_", " ").replace("-", " ").title()
